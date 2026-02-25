"""File processing service for extracting text from various file types."""

import io
from pathlib import Path

import chardet


class FileProcessor:
    """Processes uploaded files and extracts text content."""

    SUPPORTED_EXTENSIONS = {
        ".txt": "text",
        ".md": "text",
        ".csv": "text",
        ".json": "text",
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".xlsx": "excel",
        ".xls": "excel",
    }

    async def process_file(self, filename: str, content: bytes) -> dict:
        """Process a file and extract its text content."""
        ext = Path(filename).suffix.lower()
        file_type = self.SUPPORTED_EXTENSIONS.get(ext, "unknown")

        if file_type == "text":
            text = self._extract_text(content)
        elif file_type == "pdf":
            text = self._extract_pdf(content)
        elif file_type == "docx":
            text = self._extract_docx(content)
        elif file_type == "excel":
            text = self._extract_excel(content)
        else:
            # Try to extract as text anyway
            text = self._extract_text(content)

        return {
            "filename": filename,
            "file_type": file_type,
            "text": text,
            "char_count": len(text),
        }

    def _extract_text(self, content: bytes) -> str:
        """Extract text from plain text files with encoding detection."""
        try:
            # Detect encoding
            detected = chardet.detect(content)
            encoding = detected.get("encoding", "utf-8") or "utf-8"
            return content.decode(encoding)
        except Exception:
            # Fallback to utf-8 with error handling
            return content.decode("utf-8", errors="ignore")

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF files."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            texts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
            return "\n\n".join(texts)
        except Exception as e:
            return f"[PDF読み取りエラー: {str(e)}]"

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from Word documents."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        texts.append(" | ".join(row_texts))
            return "\n".join(texts)
        except Exception as e:
            return f"[Word読み取りエラー: {str(e)}]"

    def _extract_excel(self, content: bytes) -> str:
        """Extract text from Excel files."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(io.BytesIO(content), data_only=True)
            texts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                texts.append(f"## シート: {sheet_name}")
                for row in sheet.iter_rows(values_only=True):
                    row_texts = [str(cell) for cell in row if cell is not None]
                    if row_texts:
                        texts.append(" | ".join(row_texts))
            return "\n".join(texts)
        except Exception as e:
            return f"[Excel読み取りエラー: {str(e)}]"

    async def summarize_files(self, files_data: list[dict], llm_service) -> str:
        """Use LLM to summarize multiple files into a coherent brief."""
        if not files_data:
            return ""

        # Combine all file contents
        combined_text = ""
        for file_data in files_data:
            combined_text += f"\n\n--- ファイル: {file_data['filename']} ---\n"
            # Truncate very long files
            text = file_data["text"]
            if len(text) > 10000:
                text = text[:10000] + "\n...[以下省略]..."
            combined_text += text

        # Use LLM to analyze and extract key information
        system_prompt = """あなたはマーケティング戦略の専門家です。
添付されたファイルの内容を分析し、戦略プランニングに必要な情報を抽出してください。

以下の観点で情報を整理してください：
1. 製品・サービスの特徴
2. ターゲット顧客情報
3. 市場・競合情報
4. 現状の課題
5. 目標・KPI
6. その他の重要な情報

簡潔かつ構造的にまとめてください。"""

        user_prompt = f"""以下のファイル内容を分析し、戦略プランニングに必要な情報を抽出してください：

{combined_text}"""

        try:
            summary = await llm_service.generate(
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                temperature=0.5,
                max_tokens=2000,
            )
            return summary
        except Exception as e:
            return f"[ファイル分析エラー: {str(e)}]"


# Singleton instance
file_processor = FileProcessor()
